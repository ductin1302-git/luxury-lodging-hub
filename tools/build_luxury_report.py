from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT_DIR = Path(r"D:\Luxury_Lodging_Hub_Report")
ASSET_DIR = OUT_DIR / "assets"
OUTPUT_DOCX = Path(r"D:\Bao_cao_Luxury_Lodging_Hub_chuan_chinh.docx")
OLD_DOCX = Path(r"C:\Users\ACER\Downloads\Bao_cao_Luxury_Lodging_Hub_rut_gon_34_trang.docx")

BLUE = "1F4E79"
DARK = "0B2545"
LIGHT_BLUE = "EAF3F8"
LIGHT_GRAY = "F4F6F8"
MID_GRAY = "D9E2EC"
TEXT = "222222"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="B8C7D9", size="6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Times New Roman", size=None, bold=None, color=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "Nhấn Ctrl+A rồi F9 để cập nhật mục lục nếu Word chưa tự cập nhật."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 8),
        ("Heading 2", 14, BLUE, 10, 6),
        ("Heading 3", 13, DARK, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    if "Caption VN" not in [s.name for s in styles]:
        st = styles.add_style("Caption VN", 1)
    else:
        st = styles["Caption VN"]
    st.font.name = "Times New Roman"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    st.font.size = Pt(11)
    st.font.italic = True
    st.font.color.rgb = RGBColor.from_string("555555")
    st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_after = Pt(8)


def add_para(doc: Document, text: str = "", style: str | None = None, align=None, bold=False) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    if style is None:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table)
    set_cell_margins(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, cell in enumerate(hdr.cells):
        cell.text = headers[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_border(cell)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_run_font(r, size=11, bold=True, color=DARK)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    set_run_font(r, size=11)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)


def font(size=34, bold=False):
    path = r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"
    return ImageFont.truetype(path, size)


def rounded_box(draw, xy, fill, outline, text, fnt, text_fill=f"#{DARK}", radius=22, align="center") -> None:
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    lines = text.split("\n")
    total_h = sum(draw.textbbox((0, 0), line, font=fnt)[3] for line in lines) + (len(lines) - 1) * 8
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        w = bbox[2] - bbox[0]
        x = x1 + 18 if align == "left" else x1 + (x2 - x1 - w) / 2
        draw.text((x, y), line, font=fnt, fill=text_fill)
        y += (bbox[3] - bbox[1]) + 8


def arrow(draw, start, end, color=f"#{BLUE}", width=5) -> None:
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    import math

    ang = math.atan2(y2 - y1, x2 - x1)
    size = 18
    pts = [
        (x2, y2),
        (x2 - size * math.cos(ang - math.pi / 7), y2 - size * math.sin(ang - math.pi / 7)),
        (x2 - size * math.cos(ang + math.pi / 7), y2 - size * math.sin(ang + math.pi / 7)),
    ]
    draw.polygon(pts, fill=color)


def save_diagram(name: str, title: str, boxes: list[tuple[int, int, int, int, str]], arrows: list[tuple[tuple[int, int], tuple[int, int]]]) -> Path:
    img = Image.new("RGB", (1800, 1050), "white")
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, 1800, 1050), fill="#FFFFFF")
    d.rectangle((0, 0, 1800, 110), fill="#0B2545")
    d.text((60, 30), title, font=font(40, True), fill="white")
    for a, b in arrows:
        arrow(d, a, b)
    colors = ["#EAF3F8", "#F4F6F8", "#E8F5E9", "#FFF8E1", "#FDECEC"]
    for idx, (x1, y1, x2, y2, txt) in enumerate(boxes):
        rounded_box(d, (x1, y1, x2, y2), colors[idx % len(colors)], "#6C8EBF", txt, font(28, True))
    out = ASSET_DIR / f"{name}.png"
    img.save(out, quality=95)
    return out


def extract_old_images() -> list[Path]:
    import zipfile

    if not OLD_DOCX.exists():
        return []
    target = ASSET_DIR / "old_media"
    target.mkdir(parents=True, exist_ok=True)
    paths: list[Path] = []
    with zipfile.ZipFile(OLD_DOCX) as z:
        for idx, name in enumerate([n for n in z.namelist() if n.startswith("word/media/")], 1):
            suffix = Path(name).suffix.lower() or ".png"
            out = target / f"old_{idx:02d}{suffix}"
            out.write_bytes(z.read(name))
            paths.append(out)
    return paths


def create_contact_sheet(name: str, title: str, images: list[Path]) -> Path | None:
    usable = []
    for p in images:
        try:
            img = Image.open(p).convert("RGB")
            if img.width >= 300 and img.height >= 180:
                usable.append(img.copy())
        except Exception:
            continue
    if not usable:
        return None
    usable = usable[:4]
    canvas = Image.new("RGB", (1800, 1200), "white")
    d = ImageDraw.Draw(canvas)
    d.rectangle((0, 0, 1800, 115), fill="#0B2545")
    d.text((60, 33), title, font=font(38, True), fill="white")
    slots = [(80, 170, 860, 560), (940, 170, 1720, 560), (80, 660, 860, 1050), (940, 660, 1720, 1050)]
    labels = ["(a)", "(b)", "(c)", "(d)"]
    for img, slot, label in zip(usable, slots, labels):
        x1, y1, x2, y2 = slot
        box_w, box_h = x2 - x1, y2 - y1
        img.thumbnail((box_w, box_h - 42))
        px = x1 + (box_w - img.width) // 2
        py = y1 + 32 + (box_h - 42 - img.height) // 2
        d.rounded_rectangle((x1, y1, x2, y2), radius=18, outline="#B8C7D9", width=3, fill="#F8FAFC")
        d.text((x1 + 18, y1 + 10), label, font=font(22, True), fill="#0B2545")
        canvas.paste(img, (px, py))
    out = ASSET_DIR / f"{name}.png"
    canvas.save(out, quality=95)
    return out


def create_diagrams() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = {}
    diagrams["architecture"] = save_diagram(
        "architecture",
        "Kiến trúc tổng thể Luxury Lodging Hub",
        [
            (80, 190, 390, 340, "Người dùng\nWeb/Mobile"),
            (80, 520, 390, 670, "Quản trị viên\nAdmin Portal"),
            (600, 190, 950, 340, "Frontend\nReact + Vite"),
            (600, 520, 950, 670, "Admin UI\nReact Router"),
            (1120, 330, 1480, 520, "Backend API\nNestJS + Prisma"),
            (1510, 190, 1730, 340, "PostgreSQL\nAiven/Local"),
            (1510, 520, 1730, 670, "Dịch vụ ngoài\nMoMo, Google,\nEmail, OpenAI"),
        ],
        [
            ((390, 265), (600, 265)),
            ((390, 595), (600, 595)),
            ((950, 265), (1120, 400)),
            ((950, 595), (1120, 465)),
            ((1480, 400), (1510, 265)),
            ((1480, 465), (1510, 595)),
        ],
    )
    diagrams["booking_flow"] = save_diagram(
        "booking_flow",
        "Luồng nghiệp vụ đặt phòng và thanh toán",
        [
            (70, 240, 330, 390, "Tìm kiếm\nkhách sạn"),
            (420, 240, 680, 390, "Chọn phòng\nngày lưu trú"),
            (770, 240, 1030, 390, "Tạo booking\nkiểm tra tồn"),
            (1120, 240, 1380, 390, "Thanh toán\nMoMo/khác"),
            (1470, 240, 1730, 390, "Xác nhận\nthông báo"),
            (770, 560, 1030, 720, "Lịch sử booking\nhủy/đánh giá"),
        ],
        [
            ((330, 315), (420, 315)),
            ((680, 315), (770, 315)),
            ((1030, 315), (1120, 315)),
            ((1380, 315), (1470, 315)),
            ((900, 390), (900, 560)),
        ],
    )
    diagrams["erd"] = save_diagram(
        "erd",
        "Mô hình dữ liệu rút gọn",
        [
            (80, 180, 430, 340, "app_users\nid, email, role"),
            (570, 180, 920, 340, "bookings\ncode, dates, status"),
            (1060, 180, 1410, 340, "booking_items\nroom, nights, subtotal"),
            (570, 520, 920, 680, "hotels\nname, city, rating"),
            (1060, 520, 1410, 680, "rooms\nprice, capacity, stock"),
            (80, 520, 430, 680, "reviews / contacts\nfeedback, support"),
            (1460, 350, 1740, 510, "payments\nmethod, amount,\ntransaction"),
        ],
        [
            ((430, 260), (570, 260)),
            ((920, 260), (1060, 260)),
            ((745, 340), (745, 520)),
            ((920, 600), (1060, 600)),
            ((920, 300), (1460, 410)),
            ((430, 600), (570, 600)),
        ],
    )
    diagrams["deployment"] = save_diagram(
        "deployment",
        "Mô hình triển khai và vận hành",
        [
            (90, 220, 430, 380, "Vercel\nFrontend SPA"),
            (560, 220, 900, 380, "Render/Node\nNestJS API"),
            (1030, 220, 1370, 380, "PostgreSQL\nAiven hoặc local"),
            (1500, 220, 1740, 380, "pgAdmin\nQuản trị CSDL"),
            (560, 570, 900, 730, "GitHub\nSource control"),
            (1030, 570, 1370, 730, "Backup/Logs\nGiám sát lỗi"),
        ],
        [
            ((430, 300), (560, 300)),
            ((900, 300), (1030, 300)),
            ((1370, 300), (1500, 300)),
            ((730, 570), (730, 380)),
            ((1200, 570), (1200, 380)),
        ],
    )
    old_images = extract_old_images()
    # The original report stores diagrams first, then UI screenshots. Skip the first few media files
    # and reuse the remaining images as professional contact sheets.
    screenshot_pool = old_images[6:] if len(old_images) > 10 else old_images
    sheet_specs = [
        ("ui_user", "Minh họa giao diện người dùng", screenshot_pool[0:4]),
        ("ui_auth_booking", "Minh họa xác thực, đặt phòng và thanh toán", screenshot_pool[4:8]),
        ("ui_admin", "Minh họa giao diện quản trị", screenshot_pool[8:12]),
        ("ui_management", "Minh họa quản lý nghiệp vụ", screenshot_pool[12:16]),
    ]
    for key, title, imgs in sheet_specs:
        sheet = create_contact_sheet(key, title, imgs)
        if sheet:
            diagrams[key] = sheet
    return diagrams


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Caption VN")
    p.add_run(text)


def add_figure(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.25))
    add_caption(doc, caption)


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        section.different_first_page_header_footer = True
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = ""
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run("Luxury Lodging Hub - Báo cáo đồ án")
        set_run_font(r, size=9, color="666666")
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = ""
        add_page_number(fp)


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = create_diagrams()
    doc = Document()
    style_document(doc)

    # Cover page
    for text, size, bold, space in [
        ("BỘ CÔNG THƯƠNG", 13, True, 0),
        ("TRƯỜNG CAO ĐẲNG CÔNG THƯƠNG TP. HỒ CHÍ MINH", 13, True, 0),
        ("KHOA CÔNG NGHỆ THÔNG TIN", 13, True, 18),
        ("BÁO CÁO ĐỒ ÁN TỐT NGHIỆP", 20, True, 22),
        ("ĐỀ TÀI", 14, True, 8),
        ("XÂY DỰNG HỆ THỐNG WEBSITE", 17, True, 2),
        ("ĐẶT PHÒNG KHÁCH SẠN", 17, True, 2),
        ("LUXURY LODGING HUB", 22, True, 24),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space)
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=DARK if "LUXURY" in text else TEXT)

    cover_lines = [
        ("Giảng viên hướng dẫn:", "Trần Nhật Nam"),
        ("Sinh viên thực hiện:", "Phan Thanh Đức Tín"),
        ("Mã số sinh viên:", "2123110173"),
        ("Lớp:", "CCQ2311E"),
        ("Niên khóa:", "2023 - 2026"),
    ]
    table = doc.add_table(rows=len(cover_lines), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (a, b) in enumerate(cover_lines):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
        for c in table.rows[i].cells:
            for p in c.paragraphs:
                for r in p.runs:
                    set_run_font(r, size=13, bold=(c == table.rows[i].cells[0]))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    r = p.add_run("TP. Hồ Chí Minh - Tháng 07/2026")
    set_run_font(r, size=13, bold=True)

    doc.add_page_break()

    add_para(doc, "LỜI NHẬN XÉT CỦA GIẢNG VIÊN", "Heading 1", WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(10):
        add_para(doc, "................................................................................................................................................")
    add_para(doc, "TP. Hồ Chí Minh, ngày …… tháng …… năm 2026", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_para(doc, "Giảng viên hướng dẫn", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
    doc.add_page_break()

    add_para(doc, "LỜI CẢM ƠN", "Heading 1", WD_ALIGN_PARAGRAPH.CENTER)
    for text in [
        "Em xin trân trọng cảm ơn thầy Trần Nhật Nam đã hướng dẫn, góp ý và hỗ trợ em trong quá trình thực hiện đề tài “Xây dựng hệ thống website đặt phòng khách sạn Luxury Lodging Hub”. Những góp ý của thầy giúp em hoàn thiện cách phân tích yêu cầu, thiết kế hệ thống, tổ chức mã nguồn và trình bày báo cáo theo hướng chặt chẽ hơn.",
        "Em xin cảm ơn quý thầy cô Khoa Công nghệ Thông tin, Trường Cao đẳng Công Thương TP. Hồ Chí Minh đã tạo nền tảng kiến thức về lập trình web, cơ sở dữ liệu, phân tích thiết kế hệ thống và quy trình triển khai phần mềm. Đây là cơ sở quan trọng để em xây dựng sản phẩm có tính ứng dụng thực tế.",
        "Do thời gian và kinh nghiệm thực tế còn hạn chế, sản phẩm vẫn có những điểm cần tiếp tục cải tiến. Em mong nhận được ý kiến đóng góp của quý thầy cô để hệ thống ngày càng hoàn thiện hơn.",
    ]:
        add_para(doc, text)
    add_para(doc, "Em xin chân thành cảm ơn!", align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()

    add_para(doc, "MỤC LỤC", "Heading 1", WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    add_field(p, r'TOC \o "1-3" \h \z \u')
    doc.add_page_break()

    add_para(doc, "DANH MỤC HÌNH", "Heading 1", WD_ALIGN_PARAGRAPH.CENTER)
    figures = [
        "Hình 1. Kiến trúc tổng thể Luxury Lodging Hub",
        "Hình 2. Luồng nghiệp vụ đặt phòng và thanh toán",
        "Hình 3. Mô hình dữ liệu rút gọn",
        "Hình 4. Nhóm giao diện phía người dùng",
        "Hình 5. Nhóm giao diện xác thực, đặt phòng và thanh toán",
        "Hình 6. Nhóm giao diện quản trị",
        "Hình 7. Nhóm giao diện quản lý nghiệp vụ",
        "Hình 8. Mô hình triển khai và vận hành",
    ]
    for f in figures:
        add_para(doc, f)
    add_para(doc, "DANH MỤC BẢNG", "Heading 1", WD_ALIGN_PARAGRAPH.CENTER)
    tables = [
        "Bảng 1. Danh mục chữ viết tắt",
        "Bảng 2. Công nghệ sử dụng trong dự án",
        "Bảng 3. Actor và phạm vi quyền",
        "Bảng 4. Danh sách chức năng hệ thống",
        "Bảng 5. Nhóm bảng dữ liệu nghiệp vụ",
        "Bảng 6. API Backend chính",
        "Bảng 7. Từ điển dữ liệu chính",
        "Bảng 8. Kịch bản kiểm thử chức năng",
        "Bảng 9. Checklist triển khai và vận hành",
        "Bảng 10. Ma trận đánh giá chất lượng",
        "Bảng 11. Cấu hình môi trường quan trọng",
    ]
    for t in tables:
        add_para(doc, t)
    doc.add_page_break()

    add_para(doc, "DANH MỤC CHỮ VIẾT TẮT", "Heading 1", WD_ALIGN_PARAGRAPH.CENTER)
    add_table(doc, ["Từ viết tắt", "Diễn giải"], [
        ["API", "Application Programming Interface - giao diện lập trình ứng dụng"],
        ["DTO", "Data Transfer Object - đối tượng truyền dữ liệu giữa client và server"],
        ["JWT", "JSON Web Token - cơ chế xác thực phiên đăng nhập"],
        ["ORM", "Object-Relational Mapping - ánh xạ đối tượng và cơ sở dữ liệu"],
        ["OTP", "One-Time Password - mã xác thực dùng một lần"],
        ["REST", "Representational State Transfer - phong cách thiết kế API"],
        ["SPA", "Single Page Application - ứng dụng web một trang"],
        ["UI/UX", "Giao diện người dùng và trải nghiệm người dùng"],
    ], [1.6, 4.9])
    add_caption(doc, "Bảng 1. Danh mục chữ viết tắt")
    doc.add_page_break()

    add_para(doc, "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI", "Heading 1")
    add_para(doc, "1.1. Lý do chọn đề tài", "Heading 2")
    add_para(doc, "Du lịch và dịch vụ lưu trú đang chuyển dịch mạnh sang mô hình số hóa. Người dùng có xu hướng tìm kiếm khách sạn, so sánh phòng, xem đánh giá, đặt phòng và thanh toán trực tuyến trong cùng một hệ thống. Đối với đơn vị vận hành khách sạn, một nền tảng quản trị tập trung giúp giảm thao tác thủ công, hạn chế sai sót dữ liệu, theo dõi booking và chăm sóc khách hàng hiệu quả hơn.")
    add_para(doc, "Luxury Lodging Hub được xây dựng nhằm mô phỏng một hệ thống đặt phòng khách sạn hiện đại, có đầy đủ các phân hệ cho khách hàng và quản trị viên. Đề tài phù hợp để vận dụng kiến thức về React, NestJS, PostgreSQL, Prisma ORM, xác thực JWT, tích hợp Google, thanh toán MoMo, upload tệp, thông báo và hỗ trợ AI.")
    add_para(doc, "1.2. Mục tiêu đề tài", "Heading 2")
    add_bullets(doc, [
        "Xây dựng website đặt phòng khách sạn có khả năng tìm kiếm, lọc, xem chi tiết khách sạn/phòng và đặt phòng trực tuyến.",
        "Xây dựng trang quản trị để quản lý khách sạn, phòng, booking, khách hàng, khuyến mãi, đánh giá, liên hệ, tin tức, thông báo và thống kê.",
        "Tích hợp xác thực tài khoản bằng email/mật khẩu, OTP, Google OAuth và phân quyền theo vai trò.",
        "Tích hợp thanh toán MoMo, ghi nhận trạng thái giao dịch, cập nhật booking và gửi thông báo tương ứng.",
        "Thiết kế cơ sở dữ liệu PostgreSQL đảm bảo ràng buộc nghiệp vụ, truy vấn ổn định và dễ mở rộng.",
    ])
    add_para(doc, "1.3. Phạm vi thực hiện", "Heading 2")
    add_para(doc, "Phạm vi hệ thống tập trung vào nghiệp vụ đặt phòng khách sạn trong môi trường phát triển và thử nghiệm. Các chức năng liên quan đến đối soát thanh toán thực tế, hóa đơn điện tử, hoàn tiền tự động đa cổng, kiểm thử tải quy mô lớn và vận hành production nhiều vùng địa lý được đưa vào hướng phát triển.")
    add_para(doc, "1.4. Phương pháp thực hiện", "Heading 2")
    add_numbered(doc, [
        "Khảo sát nghiệp vụ đặt phòng khách sạn và phân rã yêu cầu theo actor.",
        "Thiết kế use case, luồng nghiệp vụ, mô hình dữ liệu và API.",
        "Xây dựng frontend React/TypeScript và backend NestJS theo kiến trúc REST API.",
        "Tích hợp PostgreSQL thông qua Prisma ORM, kiểm tra dữ liệu bằng pgAdmin.",
        "Kiểm thử thủ công theo kịch bản nghiệp vụ, kiểm tra build và rà soát lỗi hiển thị tiếng Việt.",
    ])

    add_para(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ", "Heading 1")
    add_para(doc, "2.1. Kiến trúc ứng dụng web", "Heading 2")
    add_para(doc, "Hệ thống được tổ chức theo mô hình client-server. Frontend là SPA chạy trên trình duyệt, giao tiếp với backend thông qua REST API. Backend xử lý xác thực, phân quyền, nghiệp vụ đặt phòng, thanh toán, upload và truy vấn cơ sở dữ liệu. Cách tổ chức này giúp tách biệt giao diện với nghiệp vụ, thuận lợi cho mở rộng và triển khai độc lập.")
    add_para(doc, "2.2. Công nghệ sử dụng", "Heading 2")
    add_table(doc, ["Công nghệ", "Vai trò", "Lý do lựa chọn"], [
        ["React + TypeScript", "Xây dựng giao diện người dùng", "Hỗ trợ component hóa, type safety, dễ bảo trì"],
        ["Vite", "Công cụ build frontend", "Khởi động nhanh, cấu hình gọn, tối ưu bundle"],
        ["Tailwind CSS + shadcn/ui", "Thiết kế giao diện", "Tạo giao diện nhất quán, hiện đại, dễ responsive"],
        ["NestJS", "Backend REST API", "Kiến trúc module rõ ràng, hỗ trợ guard, controller, service"],
        ["Prisma ORM", "Truy cập dữ liệu", "Model rõ ràng, type-safe, thuận tiện truy vấn PostgreSQL"],
        ["PostgreSQL", "Cơ sở dữ liệu quan hệ", "Hỗ trợ ràng buộc, transaction, dữ liệu phức tạp"],
        ["JWT + Passport", "Xác thực và phân quyền", "Phù hợp SPA/API, dễ bảo vệ endpoint"],
        ["MoMo Sandbox", "Thanh toán trực tuyến", "Mô phỏng luồng thanh toán thực tế tại Việt Nam"],
        ["Google OAuth", "Đăng nhập nhanh", "Tăng trải nghiệm người dùng, giảm ma sát đăng ký"],
        ["OpenAI API", "AI hỗ trợ khách hàng", "Trả lời câu hỏi, tư vấn theo phạm vi dữ liệu hệ thống"],
    ], [1.45, 1.95, 3.1])
    add_caption(doc, "Bảng 2. Công nghệ sử dụng trong dự án")
    add_para(doc, "2.3. Bảo mật và tích hợp", "Heading 2")
    add_para(doc, "Các API quan trọng được bảo vệ bằng JWT Guard và phân quyền role admin/staff/customer. Dữ liệu mật khẩu được băm trước khi lưu. OTP email/điện thoại được dùng cho các luồng xác thực. Các tích hợp bên ngoài như MoMo, Google, SMTP và OpenAI cần được cấu hình bằng biến môi trường, không hard-code trong mã nguồn khi triển khai thực tế.")

    add_para(doc, "CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", "Heading 1")
    add_para(doc, "3.1. Actor và quyền hạn", "Heading 2")
    add_table(doc, ["STT", "Actor", "Quyền và vai trò"], [
        ["1", "Khách vãng lai", "Xem trang chủ, danh sách khách sạn, chi tiết phòng, tin tức, khuyến mãi và gửi liên hệ."],
        ["2", "Khách hàng", "Đăng nhập, đặt phòng, thanh toán, theo dõi booking, hủy theo điều kiện, đánh giá, yêu thích, nhận thông báo."],
        ["3", "Nhân viên", "Theo dõi booking, phản hồi liên hệ, quản lý nội dung theo phạm vi được cấp quyền."],
        ["4", "Quản trị viên", "Quản lý toàn bộ khách sạn, phòng, người dùng, booking, khuyến mãi, đánh giá, thống kê, cấu hình."],
        ["5", "Hệ thống ngoài", "Google OAuth, MoMo, SMTP email, OpenAI API và PostgreSQL managed service."],
    ], [0.5, 1.5, 4.5])
    add_caption(doc, "Bảng 3. Actor và phạm vi quyền")
    add_para(doc, "3.2. Yêu cầu chức năng", "Heading 2")
    add_table(doc, ["Nhóm", "Chức năng chính", "Ghi chú nghiệp vụ"], [
        ["Xác thực", "Đăng ký, OTP email/phone, đăng nhập, Google OAuth, refresh token, quên/đổi/đặt mật khẩu", "Bảo vệ API bằng JWT và role guard."],
        ["Khách sạn/phòng", "Tìm kiếm, lọc, xem chi tiết, tiện ích, hình ảnh, chính sách, tồn phòng", "Kiểm tra trùng lịch trước khi tạo booking."],
        ["Đặt phòng", "Tạo booking, booking_items, tính đêm, số khách, số phòng, giảm giá, hủy booking", "Không cho đặt vượt tồn phòng."],
        ["Thanh toán", "Tạo giao dịch MoMo, nhận return/IPN, cập nhật payment_status", "Ghi transactionId, adminNote và thông báo."],
        ["Tương tác", "Yêu thích, đánh giá có ảnh/video, liên hệ hỗ trợ, thông báo", "Review chỉ hợp lệ khi booking đã hoàn tất."],
        ["Quản trị", "CRUD khách sạn, phòng, booking, khách hàng, khuyến mãi, đánh giá, liên hệ, tin tức, dashboard", "Phân quyền admin/staff."],
        ["AI hỗ trợ", "Widget tư vấn theo câu hỏi, dữ liệu khách sạn, booking khi có đăng nhập", "Giới hạn phạm vi trả lời để tránh tư vấn sai dữ liệu."],
    ], [1.25, 2.35, 2.9])
    add_caption(doc, "Bảng 4. Danh sách chức năng hệ thống")
    add_para(doc, "3.3. Kiến trúc tổng thể", "Heading 2")
    add_figure(doc, diagrams["architecture"], "Hình 1. Kiến trúc tổng thể Luxury Lodging Hub")
    add_para(doc, "Frontend và Admin Portal dùng chung nền React nhưng được phân tách bằng route và guard. Backend NestJS đóng vai trò trung tâm xử lý nghiệp vụ. Prisma ORM chuẩn hóa truy cập PostgreSQL. Các dịch vụ ngoài được kết nối thông qua module riêng để dễ thay đổi cấu hình triển khai.")
    add_para(doc, "3.4. Luồng đặt phòng và thanh toán", "Heading 2")
    add_figure(doc, diagrams["booking_flow"], "Hình 2. Luồng nghiệp vụ đặt phòng và thanh toán")
    add_para(doc, "Khi khách hàng chọn phòng và ngày lưu trú, hệ thống kiểm tra số phòng còn trống bằng cách đối chiếu các booking chưa hủy có khoảng ngày giao nhau. Nếu hợp lệ, hệ thống tạo booking, booking_items và chờ thanh toán. Khi MoMo trả kết quả thành công, booking được xác nhận, payment được cập nhật và thông báo được gửi cho khách hàng/quản trị viên.")
    add_para(doc, "3.5. Mô hình dữ liệu", "Heading 2")
    add_figure(doc, diagrams["erd"], "Hình 3. Mô hình dữ liệu rút gọn")
    add_table(doc, ["Nhóm dữ liệu", "Các bảng", "Vai trò"], [
        ["Người dùng và xác thực", "app_users, otps, notifications", "Tài khoản, vai trò, xác minh OTP và thông báo."],
        ["Khách sạn và phòng", "hotels, rooms, amenities, hotel_amenities, room_amenities, hotel_images, room_images, hotel_policies", "Thông tin lưu trú, sức chứa, giá, tồn phòng, tiện ích, hình ảnh và chính sách."],
        ["Đặt phòng và thanh toán", "bookings, booking_items, payments", "Mã booking, thời gian lưu trú, chi tiết phòng, tổng tiền và giao dịch. Quan hệ bookings-payments là 1-1 trong phiên bản hiện tại."],
        ["Tương tác và khuyến mãi", "reviews, review_media, contacts, promotions, promotion_hotels, promotion_rooms", "Đánh giá, phương tiện đánh giá, liên hệ và phạm vi áp dụng khuyến mãi."],
    ], [1.75, 2.65, 2.1])
    add_caption(doc, "Bảng 5. Nhóm bảng dữ liệu nghiệp vụ")
    add_para(doc, "3.6. API Backend chính", "Heading 2")
    add_table(doc, ["Module", "Endpoint tiêu biểu", "Mục đích"], [
        ["auth", "/api/auth/login, /register, /forgot-password, /google", "Xác thực, OTP, Google OAuth, đổi/đặt lại mật khẩu."],
        ["hotels", "/api/hotels, /api/hotels/:id, /api/hotels/:id/rooms", "Tra cứu và quản trị khách sạn/phòng."],
        ["bookings", "/api/bookings, /my-bookings, /:id/status, /:id/review", "Đặt phòng, lịch sử, hủy, cập nhật trạng thái và đánh giá."],
        ["payments", "/api/payments/momo/create, /return, /ipn", "Tạo và xử lý giao dịch MoMo."],
        ["promotions", "/api/promotions", "Quản lý khuyến mãi và phạm vi áp dụng."],
        ["contacts", "/api/contacts, /api/contacts/my", "Liên hệ hỗ trợ và phản hồi khách hàng."],
        ["notifications", "/api/notifications/user, /admin, /read-all", "Thông báo cho khách hàng và quản trị viên."],
        ["ai", "/api/ai/support", "AI hỗ trợ tư vấn theo dữ liệu hệ thống."],
    ], [1.25, 2.75, 2.5])
    add_caption(doc, "Bảng 6. API Backend chính")
    add_para(doc, "3.7. Từ điển dữ liệu chính", "Heading 2")
    add_table(doc, ["Bảng", "Trường quan trọng", "Ý nghĩa nghiệp vụ"], [
        ["app_users", "id, full_name, email, phone, password_hash, role, is_active", "Lưu thông tin tài khoản, vai trò và trạng thái hoạt động."],
        ["hotels", "id, code, slug, name, city, rating_avg, price_per_night, status", "Quản lý thông tin khách sạn, vị trí, trạng thái công bố và giá tham khảo."],
        ["rooms", "id, hotel_id, name, max_guests, price_per_night, quantity_available", "Quản lý loại phòng, sức chứa, giá và số lượng có thể bán."],
        ["bookings", "id, booking_code, user_id, hotel_id, check_in, check_out, total_amount, status", "Lưu giao dịch đặt phòng và trạng thái xử lý."],
        ["booking_items", "booking_id, room_id, rooms_count, nights, line_subtotal", "Chi tiết phòng trong mỗi booking, phục vụ tính tiền và kiểm tra tồn."],
        ["payments", "booking_id, payment_method, payment_status, amount, transaction_id", "Ghi nhận thanh toán chính của booking và mã giao dịch."],
        ["promotions", "code, discount_type, discount_value, start_date, end_date, usage_limit", "Quản lý mã khuyến mãi và điều kiện áp dụng."],
        ["reviews", "hotel_id, user_id, rating, comment, is_visible", "Lưu đánh giá đã xác thực từ khách hàng."],
        ["contacts", "name, email, subject, message, status, reply_message", "Quản lý yêu cầu hỗ trợ và phản hồi."],
        ["notifications", "recipient_user_id, scope, kind, title, message, is_read", "Thông báo theo sự kiện cho khách hàng và quản trị viên."],
    ], [1.35, 2.8, 2.35])
    add_caption(doc, "Bảng 7. Từ điển dữ liệu chính")
    add_para(doc, "3.8. Ràng buộc nghiệp vụ quan trọng", "Heading 2")
    add_bullets(doc, [
        "Email người dùng là duy nhất và được xử lý không phân biệt chữ hoa/thường bằng kiểu dữ liệu citext.",
        "Booking chỉ được tạo khi ngày trả phòng lớn hơn ngày nhận phòng và còn đủ số phòng trong khoảng ngày đã chọn.",
        "Review chỉ được phép tạo khi booking thuộc người dùng đã đăng nhập và trạng thái lưu trú đã hoàn tất.",
        "Promotion có thể giới hạn theo thời gian, số lượt dùng, khách sạn hoặc phòng cụ thể.",
        "Các API quản trị bắt buộc có JWT hợp lệ và role admin/staff.",
        "Khi thanh toán MoMo thành công, booking và payment phải được cập nhật nhất quán trong cùng luồng xử lý.",
    ])

    add_para(doc, "CHƯƠNG 4. KẾT QUẢ XÂY DỰNG HỆ THỐNG", "Heading 1")
    add_para(doc, "4.1. Phân hệ người dùng", "Heading 2")
    add_bullets(doc, [
        "Trang chủ giới thiệu dịch vụ, khách sạn nổi bật, ưu đãi, tin tức và điều hướng nhanh.",
        "Trang danh sách khách sạn hỗ trợ lọc theo địa điểm, ngày, số khách, số phòng, giá và tiêu chí liên quan.",
        "Trang chi tiết khách sạn/phòng hiển thị hình ảnh, tiện ích, chính sách, giá, số lượng, đánh giá và lựa chọn đặt phòng.",
        "Luồng đặt phòng ghi nhận thông tin khách, kiểm tra tồn phòng, áp dụng khuyến mãi và chuyển sang thanh toán.",
        "Tài khoản cá nhân có hồ sơ, lịch sử booking, chi tiết booking, hủy booking, đánh giá và danh sách yêu thích.",
        "Khách hàng có thể gửi liên hệ, xem phản hồi hỗ trợ và nhận thông báo theo sự kiện.",
    ])
    add_para(doc, "4.2. Phân hệ quản trị", "Heading 2")
    add_bullets(doc, [
        "Dashboard tổng quan hiển thị số liệu booking, doanh thu, khách hàng và trạng thái vận hành.",
        "Quản lý khách sạn/phòng cho phép thêm, sửa, xóa mềm, upload ảnh, cấu hình giá, sức chứa, tiện ích và chính sách.",
        "Quản lý booking hỗ trợ xem chi tiết, cập nhật trạng thái, theo dõi thanh toán và xử lý hủy.",
        "Quản lý khách hàng, đánh giá, khuyến mãi, liên hệ, thông báo, tin tức và cấu hình hệ thống.",
        "Phân quyền admin/staff/customer giúp hạn chế truy cập trái phép vào API và giao diện quản trị.",
    ])
    if "ui_user" in diagrams:
        add_para(doc, "4.3. Minh họa giao diện hệ thống", "Heading 2")
        add_figure(doc, diagrams["ui_user"], "Hình 4. Nhóm giao diện phía người dùng")
    if "ui_auth_booking" in diagrams:
        add_figure(doc, diagrams["ui_auth_booking"], "Hình 5. Nhóm giao diện xác thực, đặt phòng và thanh toán")
    if "ui_admin" in diagrams:
        add_figure(doc, diagrams["ui_admin"], "Hình 6. Nhóm giao diện quản trị")
    if "ui_management" in diagrams:
        add_figure(doc, diagrams["ui_management"], "Hình 7. Nhóm giao diện quản lý nghiệp vụ")
    add_para(doc, "4.4. Tích hợp thanh toán, thông báo và AI", "Heading 2")
    add_para(doc, "MoMo được tích hợp theo mô hình tạo giao dịch, chuyển hướng người dùng, nhận return/IPN và cập nhật trạng thái thanh toán. Hệ thống sinh thông báo cho khách hàng và quản trị viên sau các sự kiện tạo booking, thanh toán, xác nhận, check-in, check-out hoặc hủy. Widget AI hỗ trợ trả lời câu hỏi về khách sạn, phòng, booking và chính sách trong phạm vi dữ liệu được kiểm soát.")
    add_para(doc, "4.5. Triển khai và vận hành", "Heading 2")
    add_figure(doc, diagrams["deployment"], "Hình 8. Mô hình triển khai và vận hành")
    add_table(doc, ["Hạng mục", "Yêu cầu", "Trạng thái đề xuất"], [
        ["Frontend", "Build Vite và triển khai Vercel hoặc hosting tĩnh", "Đã cấu hình biến VITE_API_BASE_URL"],
        ["Backend", "Build NestJS, cấu hình PORT, JWT, CORS, public API URL", "Sẵn sàng chạy local/Render"],
        ["Database", "PostgreSQL local hoặc Aiven, sslmode=require khi dùng cloud", "Cần đảm bảo host DNS và mật khẩu đúng"],
        ["Storage", "Thư mục uploads hoặc dịch vụ lưu trữ ngoài", "Phù hợp môi trường phát triển"],
        ["Bảo mật", "Không commit secret, rotate key đã lộ, bật HTTPS production", "Bắt buộc trước khi triển khai thật"],
        ["Giám sát", "Log lỗi, backup DB, kiểm tra dung lượng upload", "Đưa vào hướng phát triển/vận hành"],
    ], [1.25, 3.0, 2.25])
    add_caption(doc, "Bảng 9. Checklist triển khai và vận hành")

    add_para(doc, "CHƯƠNG 5. KIỂM THỬ VÀ ĐÁNH GIÁ", "Heading 1")
    add_para(doc, "5.1. Kiểm thử chức năng", "Heading 2")
    add_table(doc, ["Mã", "Chức năng", "Dữ liệu/điều kiện", "Kết quả mong đợi", "KQ"], [
        ["TC01", "Đăng ký tài khoản", "Email hợp lệ, OTP đúng", "Tạo tài khoản và xác minh thành công", "Đạt"],
        ["TC02", "Đăng nhập", "Email/mật khẩu đúng và sai", "Trả JWT khi đúng; báo lỗi khi sai", "Đạt"],
        ["TC03", "Quên mật khẩu", "Email tồn tại, OTP hợp lệ", "Cho phép đặt lại mật khẩu", "Đạt"],
        ["TC04", "Tìm kiếm khách sạn", "Địa điểm, ngày, số khách, số phòng", "Danh sách phù hợp được trả về", "Đạt"],
        ["TC05", "Xem chi tiết phòng", "Chọn khách sạn và phòng", "Hiển thị giá, tiện ích, ảnh, tồn phòng", "Đạt"],
        ["TC06", "Tạo booking", "Ngày hợp lệ, đủ phòng", "Tạo booking_items và tổng tiền đúng", "Đạt"],
        ["TC07", "Thanh toán MoMo", "Booking chờ thanh toán", "Nhận callback và cập nhật paid/failed", "Đạt"],
        ["TC08", "Lịch sử/hủy booking", "Tài khoản đã đăng nhập", "Hiển thị đúng booking; hủy theo điều kiện", "Đạt"],
        ["TC09", "Đánh giá", "Booking đã hoàn tất", "Lưu điểm, nội dung, ảnh/video đánh giá", "Đạt"],
        ["TC10", "Liên hệ và thông báo", "Gửi yêu cầu hỗ trợ", "Admin xem, phản hồi; user nhận thông báo", "Đạt"],
        ["TC11", "Phân quyền quản trị", "Customer gọi API admin", "Từ chối truy cập; admin thao tác bình thường", "Đạt"],
        ["TC12", "Upload ảnh", "File hợp lệ/không hợp lệ", "Lưu file hợp lệ; báo lỗi rõ ràng khi sai định dạng", "Đạt"],
        ["TC13", "AI hỗ trợ", "Câu hỏi công khai và câu hỏi tài khoản", "Trả lời đúng phạm vi; yêu cầu đăng nhập khi cần", "Đạt"],
    ], [0.55, 1.55, 1.85, 2.0, 0.55])
    add_caption(doc, "Bảng 8. Kịch bản kiểm thử chức năng")
    add_para(doc, "5.2. Kiểm thử kỹ thuật", "Heading 2")
    add_bullets(doc, [
        "Backend build bằng NestJS thành công, chứng minh mã TypeScript có thể biên dịch.",
        "Frontend build bằng Vite thành công, sinh bundle production.",
        "API khách sạn và khuyến mãi trả dữ liệu đúng định dạng JSON UTF-8.",
        "Dữ liệu tiếng Việt trong PostgreSQL đã được rà soát để tránh lỗi mojibake hoặc mất dấu.",
        "Một số lỗi lint TypeScript còn tồn tại do cấu hình nghiêm ngặt với `any`; đây là nợ kỹ thuật cần xử lý ở giai đoạn hoàn thiện.",
    ])
    add_para(doc, "5.3. Đánh giá bảo mật và hiệu năng", "Heading 2")
    add_para(doc, "Hệ thống đã có nền tảng bảo mật cơ bản gồm JWT, role guard, hash mật khẩu, OTP và kiểm soát CORS. Tuy nhiên, khi triển khai thực tế cần thay toàn bộ secret đã lộ, bật HTTPS, giới hạn dung lượng upload, kiểm soát rate limit cho OTP/đăng nhập, sao lưu CSDL định kỳ và tách cấu hình production khỏi môi trường phát triển.")
    add_para(doc, "Về hiệu năng, frontend đã build được nhưng bundle chính còn lớn, nên có thể tối ưu bằng lazy loading, code splitting và tách các trang admin/client. Backend cần bổ sung cache cho dữ liệu ít thay đổi và index thêm các cột tìm kiếm khi dữ liệu tăng.")
    add_para(doc, "5.4. Ma trận chất lượng", "Heading 2")
    add_table(doc, ["Tiêu chí", "Cách đáp ứng hiện tại", "Khuyến nghị hoàn thiện"], [
        ["Tính đúng đắn", "Các luồng chính được kiểm thử thủ công và build thành công.", "Bổ sung unit test/e2e test tự động."],
        ["Bảo mật", "JWT, role guard, hash mật khẩu, OTP, CORS.", "Bật rate limit, HTTPS, rotate secret và kiểm thử OWASP."],
        ["Khả dụng", "Giao diện responsive và có thông báo trạng thái.", "Thêm trang trạng thái, retry, fallback khi dịch vụ ngoài lỗi."],
        ["Hiệu năng", "Vite build production, API tách module.", "Lazy load route, tối ưu ảnh, thêm index và cache."],
        ["Bảo trì", "NestJS module/service/controller, Prisma schema rõ ràng.", "Giảm `any`, chuẩn hóa DTO, bổ sung tài liệu API."],
        ["Vận hành", "Có biến môi trường và hướng dẫn triển khai.", "CI/CD, backup, log tập trung, giám sát uptime."],
    ], [1.35, 2.6, 2.55])
    add_caption(doc, "Bảng 10. Ma trận đánh giá chất lượng")

    add_para(doc, "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", "Heading 1")
    add_para(doc, "Kết quả đạt được", "Heading 2")
    add_para(doc, "Đề tài đã xây dựng được hệ thống website đặt phòng khách sạn với kiến trúc tương đối đầy đủ gồm frontend người dùng, trang quản trị, backend REST API, cơ sở dữ liệu PostgreSQL, xác thực, phân quyền, thanh toán MoMo, thông báo, đánh giá, liên hệ, khuyến mãi và AI hỗ trợ. Sản phẩm có khả năng chạy trong môi trường phát triển và có nền tảng để triển khai thử nghiệm thực tế.")
    add_para(doc, "Hạn chế", "Heading 2")
    add_bullets(doc, [
        "Kiểm thử tự động chưa đầy đủ; hiện chủ yếu kiểm thử thủ công và build.",
        "Một số phần tích hợp production như email thật, webhook public, lưu trữ file bền vững và giám sát log cần hoàn thiện thêm.",
        "Bundle frontend còn lớn; lint TypeScript còn nhiều cảnh báo/lỗi kiểu dữ liệu.",
        "Đối soát thanh toán, hoàn tiền tự động và hóa đơn điện tử chưa triển khai đầy đủ.",
    ])
    add_para(doc, "Hướng phát triển", "Heading 2")
    add_bullets(doc, [
        "Bổ sung test tự động cho service, controller và các luồng đặt phòng/thanh toán quan trọng.",
        "Tối ưu frontend bằng lazy loading, tách bundle admin/client và cải thiện tốc độ tải trang.",
        "Mở rộng thanh toán VNPay, ZaloPay, chuyển khoản ngân hàng và quản lý hoàn tiền.",
        "Tích hợp email/SMS tự động cho toàn bộ sự kiện booking và nhắc lịch nhận phòng.",
        "Triển khai CI/CD, backup database, monitoring, rate limiting và kiểm thử bảo mật.",
        "Nâng cấp AI hỗ trợ bằng dữ liệu có kiểm soát, lịch sử hội thoại và cơ chế chuyển cho nhân viên.",
    ])

    add_para(doc, "TÀI LIỆU THAM KHẢO", "Heading 1")
    refs = [
        "Meta Platforms, Inc., React Documentation, https://react.dev, truy cập tháng 07/2026.",
        "Microsoft, TypeScript Documentation, https://www.typescriptlang.org/docs, truy cập tháng 07/2026.",
        "Vite Team, Vite Guide, https://vite.dev, truy cập tháng 07/2026.",
        "Tailwind Labs, Tailwind CSS Documentation, https://tailwindcss.com/docs, truy cập tháng 07/2026.",
        "NestJS, NestJS Documentation, https://docs.nestjs.com, truy cập tháng 07/2026.",
        "Prisma Data, Prisma Documentation, https://www.prisma.io/docs, truy cập tháng 07/2026.",
        "PostgreSQL Global Development Group, PostgreSQL Documentation, https://www.postgresql.org/docs, truy cập tháng 07/2026.",
        "MoMo, MoMo Developer Documentation, https://developers.momo.vn, truy cập tháng 07/2026.",
        "Google for Developers, Google Identity Services, https://developers.google.com/identity, truy cập tháng 07/2026.",
        "OpenAI, OpenAI API Documentation, https://platform.openai.com/docs, truy cập tháng 07/2026.",
    ]
    add_numbered(doc, refs)

    add_para(doc, "PHỤ LỤC A. HƯỚNG DẪN CÀI ĐẶT VÀ CHẠY THỬ", "Heading 1")
    add_para(doc, "Backend:")
    add_bullets(doc, [
        "Cài thư viện: npm install",
        "Cấu hình DATABASE_URL, JWT_SECRET, FRONTEND_URL, SMTP, MoMo và OpenAI trong file .env.",
        "Build: npm run build",
        "Chạy production local: npm run start:prod",
    ])
    add_para(doc, "Frontend:")
    add_bullets(doc, [
        "Cài thư viện: npm install",
        "Cấu hình VITE_API_BASE_URL trỏ về backend.",
        "Build: npm run build",
        "Chạy dev: npm run dev",
    ])
    add_para(doc, "PHỤ LỤC B. CẤU HÌNH MÔI TRƯỜNG", "Heading 1")
    add_table(doc, ["Biến môi trường", "Mục đích", "Ghi chú"], [
        ["DATABASE_URL", "Chuỗi kết nối PostgreSQL", "Cloud PostgreSQL cần sslmode=require; local dùng database hotel_booking."],
        ["JWT_SECRET", "Khóa ký JWT", "Phải dùng chuỗi dài, không commit lên Git."],
        ["FRONTEND_URL", "Nguồn frontend được phép CORS/redirect", "Ví dụ localhost hoặc Vercel domain."],
        ["PUBLIC_API_URL", "URL backend public", "Dùng cho callback và link tích hợp."],
        ["GOOGLE_CLIENT_ID/SECRET", "Đăng nhập Google", "Cấu hình redirect URI chính xác."],
        ["MAIL_HOST/USER/PASS", "Gửi OTP và email", "Gmail cần App Password."],
        ["MOMO_*", "Thanh toán MoMo", "Sandbox và production dùng bộ key khác nhau."],
        ["OPENAI_API_KEY", "AI hỗ trợ", "Cần giới hạn quyền và theo dõi chi phí."],
    ], [1.85, 2.3, 2.35])
    add_caption(doc, "Bảng 11. Cấu hình môi trường quan trọng")
    add_para(doc, "PHỤ LỤC C. CHECKLIST NỘP ĐỒ ÁN", "Heading 1")
    add_bullets(doc, [
        "Kiểm tra lại bìa, thông tin sinh viên, giảng viên hướng dẫn và thời gian thực hiện.",
        "Cập nhật mục lục, danh mục hình, danh mục bảng sau mọi chỉnh sửa cuối.",
        "Chấp nhận toàn bộ tracked changes và xóa comment trước khi nộp.",
        "Kiểm tra file Word/PDF không lỗi font, không vỡ bảng và không mất hình.",
        "Đính kèm mã nguồn, database/schema, hướng dẫn chạy, tài khoản demo và video demo nếu giảng viên yêu cầu.",
        "Không nộp kèm file chứa secret thật; thay bằng `.env.example` hoặc che khóa nhạy cảm.",
    ])

    add_header_footer(doc)
    doc.core_properties.title = "Báo cáo đồ án tốt nghiệp - Luxury Lodging Hub"
    doc.core_properties.author = "Phan Thanh Đức Tín"
    doc.save(OUTPUT_DOCX)
    print(str(OUTPUT_DOCX))


if __name__ == "__main__":
    build_doc()
