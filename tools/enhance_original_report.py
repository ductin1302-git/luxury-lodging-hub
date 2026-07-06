from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


SRC = Path(r"C:\Users\ACER\Downloads\Bao_cao_Luxury_Lodging_Hub_rut_gon_34_trang.docx")
OUT = Path(r"D:\Bao_cao_Luxury_Lodging_Hub_day_du_chuan_bao_cao.docx")
ASSET = Path(r"D:\Luxury_Lodging_Hub_Report\assets")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="AFC3DA", sz="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), sz)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(3)
                para.paragraph_format.line_spacing = 1.1
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(11)
            if row_index == 0:
                set_cell_shading(cell, "E6F0F7")
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(14, 42, 71)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Mục lục sẽ được Word cập nhật tự động."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_caption(doc, text):
    p = doc.add_paragraph(text, style="Caption VN" if "Caption VN" in doc.styles else None)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_image(doc, path, width=5.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    return p


def add_table_data(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = str(value)
    style_table(table)
    return table


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def move_new_content_before(doc, before_para, builder):
    body = doc.element.body
    start = len(body) - 1
    builder()
    end = len(body) - 1
    new_elements = list(body[start:end])
    for element in new_elements:
        body.remove(element)
    target = before_para._p
    for element in new_elements:
        target.addprevious(element)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    for style_name in ["Normal", "Body Text VN", "List Bullet"]:
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(13)
        style.paragraph_format.line_spacing = 1.3
        style.paragraph_format.space_after = Pt(6)
        if style_name in ["Normal", "Body Text VN"]:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            style.paragraph_format.first_line_indent = Cm(0.7)

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)]:
        if name not in doc.styles:
            continue
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    if "Caption VN" in doc.styles:
        style = doc.styles["Caption VN"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12)
        style.font.italic = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    header.text = "Luxury Lodging Hub - Báo cáo đồ án"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(90, 90, 90)

    footer = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(footer)


def replace_static_toc(doc):
    paras = doc.paragraphs
    toc_para = next((p for p in paras if p.text.strip() == "MỤC LỤC"), None)
    list_para = next((p for p in paras if p.text.strip() == "DANH MỤC HÌNH, BẢNG VÀ CHỮ VIẾT TẮT"), None)
    if not toc_para or not list_para:
        return
    removing = False
    for p in list(doc.paragraphs):
        if p._p is toc_para._p:
            removing = True
            continue
        if p._p is list_para._p:
            break
        if removing:
            p._element.getparent().remove(p._element)
    toc_field = list_para.insert_paragraph_before()
    add_toc_field(toc_field)


def fix_text_and_cover(doc):
    replacements = {
        "ĐỀ CƯƠNG THỰC TẬP TỐT NGHIỆP": "BÁO CÁO ĐỒ ÁN TỐT NGHIỆP",
        "Báo cáo được tổ chức thành bốn chương. Chương 1 trình bày bối cảnh, mục tiêu và phương pháp thực hiện. Chương 2 tóm tắt cơ sở lý thuyết và công nghệ. Chương 3 phân tích yêu cầu, use case, luồng nghiệp vụ, kiến trúc và cơ sở dữ liệu. Chương 4 trình bày kết quả xây dựng, kiểm thử, đánh giá, kết luận và hướng phát triển.": (
            "Báo cáo được tổ chức thành năm chương. Chương 1 trình bày bối cảnh, mục tiêu và phương pháp thực hiện. "
            "Chương 2 tóm tắt cơ sở lý thuyết và công nghệ. Chương 3 phân tích yêu cầu, use case, luồng nghiệp vụ, "
            "kiến trúc, API và cơ sở dữ liệu. Chương 4 trình bày kết quả xây dựng các phân hệ. Chương 5 trình bày "
            "kiểm thử, đánh giá, triển khai và các điều kiện để hệ thống có thể sử dụng trong thực tế."
        ),
    }
    for p in doc.paragraphs:
        if p.text in replacements:
            p.text = replacements[p.text]
        if "TP Thủ Đức, ngày …. tháng …. năm 2024" in p.text:
            p.text = p.text.replace("2024", "2026")

    cover_overrides = {
        10: "Giảng viên hướng dẫn: Trần Nhật Nam",
        11: "Thực hiện: PHAN THANH ĐỨC TÍN",
        12: "2123110173 - CCQ2311E",
        20: "TP. Hồ Chí Minh - Tháng 05/2026",
    }
    for index, text in cover_overrides.items():
        if index < len(doc.paragraphs):
            doc.paragraphs[index].text = text

    for index, p in enumerate(doc.paragraphs[:25]):
        if not p.text.strip():
            continue
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        for run in p.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            if index in [6, 8]:
                run.bold = True
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(31, 78, 121)
            elif index <= 2:
                run.bold = True
                run.font.size = Pt(13)


def add_chapter_3_additions(doc):
    def builder():
        doc.add_paragraph("3.7. Thiết kế API và hợp đồng dữ liệu", style="Heading 2")
        doc.add_paragraph(
            "Các API được tổ chức theo nhóm tài nguyên để Frontend dễ sử dụng và Backend dễ mở rộng. "
            "Những API công khai phục vụ tra cứu thông tin; các API giao dịch yêu cầu đăng nhập; API quản trị "
            "bắt buộc có JWT hợp lệ và vai trò phù hợp. Cách tổ chức này giúp hệ thống kiểm soát quyền ở cả "
            "giao diện và máy chủ, tránh phụ thuộc vào việc ẩn nút thao tác trên Frontend."
        )
        add_caption(doc, "Bảng 6. Nhóm API chính của hệ thống")
        add_table_data(
            doc,
            [
                ["Nhóm API", "Endpoint tiêu biểu", "Vai trò nghiệp vụ", "Quyền truy cập"],
                ["Xác thực", "/auth/register, /auth/login, /auth/verify-otp, /auth/google", "Tạo tài khoản, đăng nhập, xác minh OTP và đăng nhập Google", "Công khai/khách hàng"],
                ["Khách sạn và phòng", "/hotels, /rooms, /amenities", "Tra cứu khách sạn, phòng, tiện ích, hình ảnh và chính sách", "Công khai, quản trị khi thêm/sửa/xóa"],
                ["Booking", "/bookings, /bookings/:id/cancel", "Tạo đặt phòng, xem lịch sử, cập nhật/hủy booking theo điều kiện", "Khách hàng và quản trị"],
                ["Thanh toán", "/payments/momo, /payments/ipn, /payments/return", "Tạo giao dịch, nhận IPN/return và cập nhật trạng thái thanh toán", "Khách hàng, hệ thống MoMo"],
                ["Quản trị", "/admin/*, /users, /promotions, /notifications, /news", "Vận hành dữ liệu, nội dung, liên hệ, thống kê và phân quyền", "Admin/staff"],
                ["AI hỗ trợ", "/ai/chat hoặc endpoint tương đương", "Trả lời câu hỏi hỗ trợ khách hàng trong phạm vi dữ liệu cho phép", "Khách/khách hàng"],
            ],
        )
        doc.add_paragraph("3.8. Ràng buộc nghiệp vụ và kiểm soát lỗi", style="Heading 2")
        doc.add_paragraph(
            "Để hệ thống có thể sử dụng thực tế, phần Backend cần kiểm soát các ràng buộc nghiệp vụ tại máy chủ "
            "thay vì chỉ kiểm tra ở giao diện. Các kiểm tra quan trọng gồm ngày nhận phòng nhỏ hơn ngày trả phòng, "
            "số khách không vượt quá sức chứa, số phòng đặt không vượt tồn kho, giá thanh toán được tính từ dữ liệu "
            "máy chủ và trạng thái booking chỉ được chuyển theo luồng hợp lệ."
        )
        add_caption(doc, "Bảng 7. Ràng buộc nghiệp vụ cần kiểm soát")
        add_table_data(
            doc,
            [
                ["Nghiệp vụ", "Ràng buộc", "Cách kiểm soát đề xuất"],
                ["Đăng ký/OTP", "Email không trùng, OTP còn hạn và khớp mã", "Unique email, thời hạn OTP, giới hạn số lần thử"],
                ["Đặt phòng", "Ngày hợp lệ, đủ sức chứa, còn phòng, tổng tiền đúng", "DTO validation, kiểm tra tồn kho, tính tiền ở Backend"],
                ["Thanh toán", "Không cập nhật thanh toán chỉ dựa trên redirect trình duyệt", "Xác minh chữ ký/IPN từ MoMo trước khi đổi trạng thái"],
                ["Hủy booking", "Chỉ hủy booking còn trong thời hạn và chưa hoàn tất", "Kiểm tra trạng thái và ngày nhận phòng"],
                ["Quản trị", "Tài khoản thường không gọi được API admin", "JWT Guard, Role Guard, kiểm tra quyền tại Controller/Service"],
                ["Dữ liệu hiển thị", "Không xóa cứng dữ liệu đang có booking liên quan", "Soft delete/trạng thái hiển thị, khóa ngoại và kiểm tra phụ thuộc"],
            ],
        )
        doc.add_paragraph("3.9. Mô hình triển khai đề xuất", style="Heading 2")
        doc.add_paragraph(
            "Khi đưa hệ thống ra môi trường thật, Frontend có thể triển khai trên Vercel hoặc hosting tĩnh, Backend "
            "triển khai trên Render/VPS, cơ sở dữ liệu dùng PostgreSQL local hoặc dịch vụ cloud như Aiven. Với cơ sở "
            "dữ liệu cloud, chuỗi kết nối cần khai báo sslmode=require và thông tin host/port/password phải lấy trực "
            "tiếp từ trang quản trị dịch vụ để tránh lỗi DNS hoặc sai mật khẩu."
        )
        add_image(doc, ASSET / "deployment.png", 5.8)
        add_caption(doc, "Hình 14. Mô hình triển khai và vận hành đề xuất")

    target = next(p for p in doc.paragraphs if p.text.strip() == "CHƯƠNG 4. KẾT QUẢ XÂY DỰNG HỆ THỐNG")
    move_new_content_before(doc, target, builder)


def add_chapter_5(doc):
    def builder():
        h = doc.add_paragraph("CHƯƠNG 5. KIỂM THỬ, TRIỂN KHAI VÀ ĐÁNH GIÁ THỰC TẾ", style="Heading 1")
        h.paragraph_format.page_break_before = True
        doc.add_paragraph("5.1. Mục tiêu kiểm thử", style="Heading 2")
        doc.add_paragraph(
            "Mục tiêu kiểm thử là xác nhận các luồng chính của hệ thống hoạt động đúng, dữ liệu không bị sai lệch "
            "giữa giao diện và cơ sở dữ liệu, phân quyền được bảo vệ ở Backend và các lỗi thường gặp được phản hồi "
            "rõ ràng cho người dùng. Kiểm thử được thực hiện theo hướng kết hợp thao tác giao diện, gọi API, quan sát "
            "log Backend và kiểm tra dữ liệu trong PostgreSQL."
        )
        doc.add_paragraph("5.2. Bảng kiểm thử bổ sung", style="Heading 2")
        add_caption(doc, "Bảng 8. Test case bổ sung cho các luồng quan trọng")
        add_table_data(
            doc,
            [
                ["Mã", "Nhóm chức năng", "Kịch bản kiểm thử", "Kết quả mong đợi", "Trạng thái"],
                ["TC11", "Tìm kiếm", "Tìm khách sạn theo địa điểm, số khách, số phòng", "Danh sách trả về đúng điều kiện, không lỗi tiếng Việt", "Đạt/cần xác nhận demo"],
                ["TC12", "Chi tiết phòng", "Mở khách sạn, xem phòng, tiện ích, chính sách", "Hiển thị ảnh, giá, sức chứa, mô tả và nút đặt phòng", "Đạt/cần xác nhận demo"],
                ["TC13", "Booking", "Tạo booking với ngày nhận/trả hợp lệ", "Tạo mã booking, lưu chi tiết phòng, tổng tiền đúng", "Đạt/cần xác nhận demo"],
                ["TC14", "Booking lỗi", "Ngày trả nhỏ hơn hoặc bằng ngày nhận", "Backend từ chối và Frontend hiển thị thông báo", "Cần kiểm tra thêm"],
                ["TC15", "MoMo", "Tạo thanh toán sandbox và nhận return/IPN", "Booking đổi trạng thái sau khi xác minh giao dịch", "Cần kiểm tra môi trường thật"],
                ["TC16", "Phân quyền", "Tài khoản thường gọi API quản trị", "API trả 401/403, không cho sửa dữ liệu", "Đạt/cần xác nhận bằng Postman"],
                ["TC17", "Admin CRUD", "Thêm/sửa/ẩn khách sạn, phòng, banner, tin tức", "Dữ liệu cập nhật đúng, giao diện reload không mất tiếng Việt", "Đạt/cần xác nhận demo"],
                ["TC18", "Thông báo/email", "Gửi OTP hoặc thông báo booking", "Email/thông báo sinh đúng nội dung tiếng Việt", "Cần cấu hình SMTP"],
                ["TC19", "AI hỗ trợ", "Hỏi thông tin phòng/chính sách", "AI trả lời trong phạm vi dữ liệu cho phép", "Đạt mức mô phỏng"],
                ["TC20", "Responsive", "Mở trang trên desktop/mobile", "Không vỡ layout, ảnh và nút thao tác hiển thị rõ", "Cần test nhiều thiết bị"],
            ],
        )
        doc.add_paragraph("5.3. Kiểm thử kỹ thuật và chất lượng mã nguồn", style="Heading 2")
        doc.add_paragraph(
            "Ngoài kiểm thử chức năng, dự án cần kiểm tra khả năng build, cấu hình biến môi trường và kết nối cơ sở "
            "dữ liệu. Backend NestJS cần build thành công trước khi triển khai; Frontend Vite cần build production để "
            "kiểm tra lỗi TypeScript, import và cấu hình API base URL. Các cảnh báo lint hoặc kiểu dữ liệu any nên "
            "được xử lý dần để mã nguồn dễ bảo trì hơn khi mở rộng."
        )
        add_caption(doc, "Bảng 9. Checklist kỹ thuật trước khi demo hoặc triển khai")
        add_table_data(
            doc,
            [
                ["Hạng mục", "Yêu cầu", "Ghi chú"],
                ["Backend build", "npm run build chạy thành công", "Kiểm tra module, DTO, service và Prisma Client"],
                ["Frontend build", "npm run build chạy thành công", "Kiểm tra import, route, biến VITE_API_BASE_URL"],
                ["Database", "PostgreSQL có schema và dữ liệu seed cần thiết", "Dùng Prisma migrate/seed hoặc import SQL"],
                ["Kết nối Aiven", "Host, port, user, password đúng; sslmode=require", "Lỗi getaddrinfo thường do sai host/DNS hoặc service đổi endpoint"],
                ["Secrets", "JWT_SECRET, SMTP, MoMo, OpenAI không commit lên Git", "Nên dùng .env.example để mô tả biến"],
                ["Unicode tiếng Việt", "Dữ liệu SQL và response API dùng UTF-8", "Kiểm tra font hiển thị và dữ liệu seed"],
            ],
        )
        doc.add_paragraph("5.4. Hướng dẫn triển khai và chạy thử", style="Heading 2")
        doc.add_paragraph(
            "Quy trình chạy thử đề xuất gồm bốn bước: cài thư viện cho Backend và Frontend, cấu hình file .env, "
            "chuẩn bị PostgreSQL, sau đó chạy ứng dụng và kiểm thử luồng chính. Khi dùng local database, DATABASE_URL "
            "có dạng postgresql://postgres:matkhau@localhost:5432/hotel_booking?schema=public. Khi dùng Aiven, chuỗi "
            "kết nối cần đúng host, port, user, password và thêm sslmode=require."
        )
        add_bullet(doc, "Backend: cài thư viện, cấu hình DATABASE_URL/JWT/SMTP/MoMo/OpenAI, chạy migration/seed, sau đó chạy npm run start:dev hoặc npm run start:prod.")
        add_bullet(doc, "Frontend: cấu hình VITE_API_BASE_URL trỏ về Backend, chạy npm run dev để demo local hoặc npm run build để kiểm tra bản production.")
        add_bullet(doc, "pgAdmin: tạo server mới, nhập Host name/address đúng từ Aiven, Port, Maintenance database, Username, Password; tại tab Parameters thêm SSL mode = require.")
        add_bullet(doc, "Sau khi kết nối DB, kiểm tra bảng hotels, rooms, bookings, payments, banners và notifications để đảm bảo dữ liệu tiếng Việt không lỗi font.")
        doc.add_paragraph("5.5. Đánh giá mức độ sẵn sàng sử dụng thực tế", style="Heading 2")
        doc.add_paragraph(
            "Với phạm vi đồ án, hệ thống đã có đủ nền tảng để mô phỏng quy trình đặt phòng trực tuyến gồm tra cứu, "
            "đặt phòng, thanh toán, quản lý booking và quản trị dữ liệu. Để sử dụng thương mại thật, hệ thống cần "
            "hoàn thiện kiểm thử tải, xử lý đặt trùng, đối soát thanh toán, sao lưu dữ liệu, giám sát log, phân quyền "
            "chi tiết và quy trình bảo vệ khóa bí mật."
        )

    target = next(p for p in doc.paragraphs if p.text.strip() == "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")
    move_new_content_before(doc, target, builder)


def add_appendices(doc):
    p = doc.add_paragraph("PHỤ LỤC A. DANH SÁCH BIẾN MÔI TRƯỜNG QUAN TRỌNG", style="Heading 1")
    p.paragraph_format.page_break_before = True
    doc.add_paragraph(
        "Phụ lục này giúp người đọc cấu hình dự án khi chạy thử. Các giá trị thật cần được lưu trong file .env cục bộ "
        "hoặc biến môi trường của nền tảng triển khai, không đưa vào báo cáo công khai hoặc repository."
    )
    add_caption(doc, "Bảng 10. Biến môi trường quan trọng")
    add_table_data(
        doc,
        [
            ["Biến", "Mục đích", "Lưu ý"],
            ["DATABASE_URL", "Chuỗi kết nối PostgreSQL", "Cloud PostgreSQL dùng sslmode=require"],
            ["JWT_SECRET", "Khóa ký token đăng nhập", "Dùng chuỗi dài, không dùng giá trị mặc định"],
            ["FRONTEND_URL", "Domain Frontend được phép redirect/CORS", "Ví dụ localhost hoặc Vercel domain"],
            ["PUBLIC_API_URL", "URL Backend public cho callback/tích hợp", "Cần đúng khi dùng MoMo/Google"],
            ["GOOGLE_CLIENT_ID/SECRET", "Đăng nhập Google", "Redirect URI phải khớp cấu hình Google"],
            ["MAIL_HOST/USER/PASS", "Gửi OTP/email", "Gmail cần App Password"],
            ["MOMO_*", "Tích hợp thanh toán MoMo", "Sandbox và production dùng key khác nhau"],
            ["OPENAI_API_KEY", "AI hỗ trợ khách hàng", "Giới hạn quyền và theo dõi chi phí"],
        ],
    )
    doc.add_paragraph("PHỤ LỤC B. CHECKLIST TRƯỚC KHI NỘP BÁO CÁO", style="Heading 1")
    for item in [
        "Cập nhật mục lục, danh mục hình, danh mục bảng sau khi chỉnh sửa lần cuối.",
        "Kiểm tra toàn bộ hình ảnh có caption, số thứ tự và không bị vỡ trang.",
        "Kiểm tra bảng không tràn lề, không mất chữ và có tiêu đề rõ ràng.",
        "Chạy thử luồng đăng ký, đăng nhập, tìm kiếm, đặt phòng, thanh toán và quản trị.",
        "Đảm bảo dữ liệu tiếng Việt trong SQL/API/Frontend hiển thị đúng UTF-8.",
        "Không nộp kèm file chứa mật khẩu thật, khóa API thật hoặc thông tin thanh toán thật.",
        "Chuẩn bị tài khoản demo, video demo và hướng dẫn chạy nếu giảng viên yêu cầu.",
    ]:
        add_bullet(doc, item)


def final_format_pass(doc):
    for table in doc.tables:
        style_table(table)
    for index, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if index <= 20:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = None
            for run in p.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            continue
        if p.style.name == "Heading 1":
            if text.startswith(("CHƯƠNG", "KẾT LUẬN", "TÀI LIỆU", "PHỤ LỤC")):
                p.paragraph_format.page_break_before = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif p.style.name in ["Normal", "Body Text VN"] and text:
            if not text.startswith(("Bảng ", "Hình ")):
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def main():
    doc = Document(str(SRC))
    configure_document(doc)
    fix_text_and_cover(doc)
    replace_static_toc(doc)
    add_chapter_3_additions(doc)
    add_chapter_5(doc)
    add_appendices(doc)
    final_format_pass(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
